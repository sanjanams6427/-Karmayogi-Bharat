"""
KB Tender RFB IN-KBL-543730-NC-RFB — Proof of Concept (POC) Packager
=====================================================================
The tender awards 50 out of 100 technical marks for a POC submission:
  - One course translated into Hindi (hin) AND Punjabi (pan)
  - Submitted as a drive link to the evaluation committee

This script bundles all required POC assets into a single ZIP and
generates a POC cover note DOCX ready for drive-link submission.

Usage:
    python scripts/poc_package.py \\
        --video course.mp4 --src eng \\
        --course-id KB_POC_001 \\
        --output ./output \\
        --agency "My Translation Agency" \\
        --contact "Dr. Priya Nair" \\
        --email "priya@agency.in"

    # If already dubbed, skip dubbing and just package existing output:
    python scripts/poc_package.py \\
        --course-id KB_POC_001 --output ./output --package-only \\
        --agency "My Translation Agency"
"""

import argparse
import json
import os
import sys
import zipfile
import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

POC_LANGS = ["hin", "pan"]   # Hindi + Punjabi — required by tender POC spec


def _dub_poc(video_path: str, src_lang: str, course_id: str,
             output_dir: str, force: bool) -> dict:
    """Run the dubbing pipeline for hin + pan and return results."""
    from pipeline.dubbing_pipeline import DubbingPipeline
    pipeline = DubbingPipeline()
    print(f"\n[POC] Dubbing {Path(video_path).name} → Hindi + Punjabi ...")
    results = pipeline.dub_course(
        video_path=video_path,
        src_lang=src_lang,
        tgt_langs=POC_LANGS,
        output_dir=output_dir,
        course_id=course_id,
        force=force,
    )
    return {lang: r for lang, r in results.items()}


def _collect_assets(output_dir: str, course_id: str) -> dict[str, list[Path]]:
    """Collect all output files for hin and pan from output_dir."""
    assets: dict[str, list[Path]] = {}
    for lang in POC_LANGS:
        lang_dir = Path(output_dir) / lang
        if not lang_dir.exists():
            lang_dir = Path(output_dir) / course_id / lang
        found = []
        for ext in (".mp4", ".mp3", ".srt", ".vtt", "_metadata.json",
                    "_qa_cert.docx", "_sync.json"):
            for p in lang_dir.glob(f"*{ext}"):
                found.append(p)
        assets[lang] = sorted(found)
    return assets


def _generate_cover_note(
    course_id: str,
    assets: dict[str, list[Path]],
    agency_name: str,
    contact_person: str,
    contact_email: str,
    output_path: str,
) -> str:
    """Generate POC cover note DOCX for evaluation committee."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading("Proof of Concept (POC) Submission", level=1)
    doc.add_heading(
        "KB iGOT Karmayogi — Language Translation & Dubbing\n"
        "RFB No.: IN-KBL-543730-NC-RFB",
        level=2,
    )
    doc.add_paragraph(f"Submission Date: {datetime.datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Submitted by: {agency_name}")
    doc.add_paragraph(f"Contact: {contact_person}  |  {contact_email}")
    doc.add_paragraph("")

    doc.add_heading("1. POC Scope", level=2)
    doc.add_paragraph(
        "This Proof of Concept demonstrates the end-to-end translation and dubbing "
        "capability of our AI pipeline for the iGOT Karmayogi platform. "
        "As required by the tender evaluation criteria (50 technical marks), "
        "one course has been translated and dubbed into:"
    )
    for lang in POC_LANGS:
        from pipeline.lang_config import LANG_NAMES
        doc.add_paragraph(
            f"  • {LANG_NAMES.get(lang, lang)} ({lang})",
            style="List Bullet",
        )
    doc.add_paragraph("")

    doc.add_heading("2. Deliverables Included", level=2)
    dt = doc.add_table(rows=1, cols=4)
    dt.style = "Table Grid"
    for i, h in enumerate(["Language", "File", "Type", "Size (KB)"]):
        dt.rows[0].cells[i].text = h

    from pipeline.lang_config import LANG_NAMES
    for lang, files in assets.items():
        lang_name = LANG_NAMES.get(lang, lang)
        if not files:
            row = dt.add_row().cells
            row[0].text = lang_name
            row[1].text = "— No output found —"
            row[2].text = ""
            row[3].text = ""
            continue
        for f in files:
            row = dt.add_row().cells
            row[0].text = lang_name
            row[1].text = f.name
            row[2].text = f.suffix.lstrip(".").upper()
            row[3].text = str(round(f.stat().st_size / 1024, 1)) if f.exists() else "—"
    doc.add_paragraph("")

    doc.add_heading("3. Technology Stack", level=2)
    for item in [
        "ASR: faster-whisper large-v3 (offline, on-premise)",
        "Translation: IndicTrans2 fine-tuned (Hindi primary) + NLLB-200 (Punjabi primary)",
        "TTS: Parler-TTS Indic Large → MMS-TTS fallback",
        "Quality scoring: heuristic + ChrF + back-translation (per segment)",
        "Content safety: regex-based filter (§3.2 — no profanity/hate speech)",
        "Trademark passthrough: registered terms preserved untranslated (§3.2)",
        "SCORM guard: non-SCORM content only (§3.1)",
        "Sovereign AI: KB_SOVEREIGN_MODE=1 — all processing on-premise (§3.2)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")

    doc.add_heading("4. Quality Summary", level=2)
    # Load metadata JSON for quality scores if available
    for lang in POC_LANGS:
        lang_name = LANG_NAMES.get(lang, lang)
        meta_files = [f for f in assets.get(lang, []) if "_metadata.json" in f.name]
        if meta_files:
            try:
                meta = json.loads(meta_files[0].read_text(encoding="utf-8"))
                qs = meta.get("quality_summary", {})
                cs = qs.get("content_safety_pass")
                doc.add_paragraph(
                    f"{lang_name}: avg_score={qs.get('avg_score', 'N/A')}  "
                    f"pass_rate={qs.get('pass_rate', 'N/A')}  "
                    f"content_safety={'✅ Pass' if cs else ('❌ Flagged' if cs is False else 'N/A')}"
                )
            except Exception:
                doc.add_paragraph(f"{lang_name}: metadata not available")
        else:
            doc.add_paragraph(f"{lang_name}: metadata not available")
    doc.add_paragraph("")

    doc.add_heading("5. Declaration", level=2)
    doc.add_paragraph(
        f"We, {agency_name}, confirm that the POC deliverables included in this "
        f"package were produced using our fully offline, on-premise AI pipeline "
        f"in compliance with all requirements of RFB IN-KBL-543730-NC-RFB. "
        f"No course content was transmitted to any external server."
    )
    doc.add_paragraph(f"Agency: {agency_name}")
    doc.add_paragraph(f"Authorised Signatory: ____________________")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("Organisation Seal: ____________________")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_poc_package(
    course_id: str,
    output_dir: str,
    agency_name: str = "Translation Agency",
    contact_person: str = "",
    contact_email: str = "",
) -> str:
    """
    Collect all POC assets, generate cover note, and bundle into a ZIP.
    Returns path to the ZIP file.
    """
    assets = _collect_assets(output_dir, course_id)
    total_files = sum(len(v) for v in assets.values())

    if total_files == 0:
        raise FileNotFoundError(
            f"No output files found for course '{course_id}' in '{output_dir}'. "
            f"Run dubbing first (remove --package-only) or check the output directory."
        )

    now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    pkg_dir = Path(output_dir) / "_poc_package"
    pkg_dir.mkdir(parents=True, exist_ok=True)

    # Generate cover note
    cover_path = str(pkg_dir / f"KB_POC_Cover_Note_{now_str}.docx")
    _generate_cover_note(
        course_id, assets, agency_name, contact_person, contact_email, cover_path
    )

    # Build ZIP
    zip_path = str(Path(output_dir) / f"KB_POC_{course_id}_{now_str}.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        # Cover note at root
        zf.write(cover_path, Path(cover_path).name)
        # Assets organised by language
        for lang, files in assets.items():
            for f in files:
                if f.exists():
                    zf.write(str(f), f"{lang}/{f.name}")

    return zip_path


def main():
    parser = argparse.ArgumentParser(
        description="KB POC Packager — bundles Hindi + Punjabi output for evaluation committee",
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--course-id", default="KB_POC_001", help="Course identifier")
    parser.add_argument("--output", default="./output", help="Output directory")
    parser.add_argument("--video", help="Source video/audio file (skip if --package-only)")
    parser.add_argument("--src", default="eng", help="Source language (default: eng)")
    parser.add_argument("--force", action="store_true", help="Force re-dub even if output exists")
    parser.add_argument("--package-only", action="store_true",
                        help="Skip dubbing — just package existing output")
    parser.add_argument("--agency", default="Translation Agency",
                        help="Agency name for cover note")
    parser.add_argument("--contact", default="", help="Contact person name")
    parser.add_argument("--email", default="", help="Contact email")
    args = parser.parse_args()

    print(f"\n{'='*60}")
    print(f"  KB POC Packager — RFB IN-KBL-543730-NC-RFB")
    print(f"  Course ID : {args.course_id}")
    print(f"  Languages : Hindi (hin) + Punjabi (pan)")
    print(f"{'='*60}\n")

    if not args.package_only:
        if not args.video:
            print("ERROR: --video is required unless --package-only is set")
            sys.exit(1)
        if not Path(args.video).exists():
            print(f"ERROR: Video file not found: {args.video}")
            sys.exit(1)
        results = _dub_poc(args.video, args.src, args.course_id, args.output, args.force)
        for lang, r in results.items():
            from pipeline.lang_config import LANG_NAMES
            status = "✅" if r.success else f"❌ {r.error}"
            print(f"  {LANG_NAMES.get(lang, lang):<15} {status}")

    print("\n[POC] Packaging assets ...")
    try:
        zip_path = build_poc_package(
            course_id=args.course_id,
            output_dir=args.output,
            agency_name=args.agency,
            contact_person=args.contact,
            contact_email=args.email,
        )
    except FileNotFoundError as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    size_mb = round(Path(zip_path).stat().st_size / 1024 / 1024, 1)
    print(f"\n{'='*60}")
    print(f"  ✅ POC package ready")
    print(f"  ZIP : {zip_path}")
    print(f"  Size: {size_mb} MB")
    print(f"  Upload this ZIP to your drive link and share with the evaluation committee.")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
