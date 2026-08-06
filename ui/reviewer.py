"""
Human Reviewer — KB Translation System
Loads pipeline metadata JSON, lets a native-language expert
approve / reject / correct each segment, and exports a signed
review certificate (DOCX) for KB submission.

No pipeline imports — pure UI logic.
"""

import json
import datetime
from pathlib import Path


# ── Load / Save ──────────────────────────────────────────────

def load_metadata(json_path: str) -> tuple[list[dict], dict]:
    """
    Parse a *_metadata.json produced by the pipeline.
    Returns (segments, raw_meta) where each segment has:
      id, start, end, source_text, translated_text, quality_score, flags
    """
    raw = json.loads(Path(json_path).read_text(encoding="utf-8"))
    transcript   = {s["id"]: s for s in raw.get("transcript", [])}
    translations = raw.get("translations", [])

    segments = []
    for t in translations:
        sid = t.get("id")  # may be int 0 — don't default to "" which causes a miss
        src = transcript.get(sid, {})
        q   = t.get("quality", {})
        segments.append({
            "id":              sid,
            "start":           src.get("start", t.get("start", 0)),
            "end":             src.get("end",   t.get("end",   0)),
            "source_text":     src.get("text",  ""),
            "translated_text": t.get("text", ""),
            "corrected_text":  t.get("text", ""),   # editable copy
            "score":           q.get("score", 1.0),
            "flags":           q.get("flags", []),
            "needs_review":    q.get("needs_review", False),
            "decision":        "",   # "approved" | "rejected" | "corrected"
        })
    return segments, raw


def review_path(json_path: str) -> Path:
    """Sidecar file: same folder, _review.json suffix."""
    p = Path(json_path)
    return p.with_name(p.stem.replace("_metadata", "") + "_review.json")


def load_review(json_path: str, segments: list[dict]) -> list[dict]:
    """Overlay any previously saved decisions onto segments (resume support)."""
    rp = review_path(json_path)
    if not rp.exists():
        return segments
    saved = {r["id"]: r for r in json.loads(rp.read_text(encoding="utf-8")).get("segments", [])}
    for seg in segments:
        if seg["id"] in saved:
            seg["decision"]       = saved[seg["id"]].get("decision", "")
            seg["corrected_text"] = saved[seg["id"]].get("corrected_text", seg["translated_text"])
    return segments


def save_review(json_path: str, segments: list[dict], reviewer: str) -> str:
    """Persist current decisions to sidecar JSON. Returns save path."""
    rp = review_path(json_path)
    payload = {
        "reviewer":   reviewer,
        "saved_at":   datetime.datetime.now().isoformat(),
        "segments":   [
            {"id": s["id"], "decision": s["decision"],
             "corrected_text": s["corrected_text"]}
            for s in segments
        ],
    }
    rp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(rp)


# ── Certificate export ────────────────────────────────────────

def export_certificate(json_path: str, segments: list[dict],
                        reviewer: str, output_path: str) -> str:
    """
    Export a signed human-review certificate (DOCX) for KB submission.
    Includes: reviewer name, date, per-segment decisions, summary stats.
    """
    from docx import Document
    from docx.shared import RGBColor

    raw      = json.loads(Path(json_path).read_text(encoding="utf-8"))
    course   = raw.get("course_id", Path(json_path).stem)
    tgt_lang = raw.get("target_lang", "")
    tgt_name = raw.get("target_lang_name", tgt_lang)
    src_lang = raw.get("source_lang", "eng")

    total     = len(segments)
    approved  = sum(1 for s in segments if s["decision"] == "approved")
    corrected = sum(1 for s in segments if s["decision"] == "corrected")
    rejected  = sum(1 for s in segments if s["decision"] == "rejected")
    pending   = total - approved - corrected - rejected

    doc = Document()
    doc.add_heading("Human Review Certificate", level=1)
    doc.add_heading("KB iGOT Karmayogi — Native Language Expert Review", level=2)
    doc.add_paragraph("")

    # Summary table
    meta_table = doc.add_table(rows=10, cols=2)
    meta_table.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Course ID",        course),
        ("Target Language",  f"{tgt_name} ({tgt_lang})"),
        ("Source Language",  src_lang),
        ("Reviewer",         reviewer),
        ("Review Date",      datetime.datetime.now().strftime("%d %B %Y %H:%M")),
        ("Total Segments",   str(total)),
        ("Approved",         str(approved)),
        ("Corrected",        str(corrected)),
        ("Rejected",         str(rejected)),
        ("Pending",          str(pending)),
    ]):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_paragraph("")
    doc.add_heading("Segment-level Review", level=2)

    seg_table = doc.add_table(rows=1, cols=5)
    seg_table.style = "Table Grid"
    for i, h in enumerate(["#", "Source", "AI Translation", "Corrected Text", "Decision"]):
        seg_table.rows[0].cells[i].text = h

    for idx, seg in enumerate(segments, 1):
        row   = seg_table.add_row().cells
        row[0].text = str(idx)
        row[1].text = seg["source_text"][:200]
        row[2].text = seg["translated_text"][:200]
        row[3].text = seg["corrected_text"][:200]
        decision    = seg["decision"] or "pending"
        row[4].text = decision
        # Colour-code decision cell
        run = row[4].paragraphs[0].runs
        if run:
            colour = {"approved": RGBColor(0, 128, 0),
                      "corrected": RGBColor(255, 140, 0),
                      "rejected":  RGBColor(200, 0, 0)}.get(decision)
            if colour:
                run[0].font.color.rgb = colour

    doc.add_paragraph("")
    doc.add_heading("Declaration", level=2)
    doc.add_paragraph(
        f"I, {reviewer}, a qualified native {tgt_name} language expert, certify that "
        f"I have reviewed all {total} translated segments for Course ID: {course}. "
        f"Approved: {approved}, Corrected: {corrected}, Rejected: {rejected}, Pending: {pending}. "
        "The content meets linguistic accuracy standards per KB RFB IN-KBL-543730-NC-RFB."
    )
    doc.add_paragraph(f"Reviewer: {reviewer}")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("Signature: ____________________")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ── Summary helpers ───────────────────────────────────────────

def review_stats(segments: list[dict]) -> str:
    total     = len(segments)
    approved  = sum(1 for s in segments if s["decision"] == "approved")
    corrected = sum(1 for s in segments if s["decision"] == "corrected")
    rejected  = sum(1 for s in segments if s["decision"] == "rejected")
    pending   = total - approved - corrected - rejected
    flagged   = sum(1 for s in segments if s["needs_review"])
    return (f"Total: {total} | ✅ Approved: {approved} | ✏️ Corrected: {corrected} | "
            f"❌ Rejected: {rejected} | ⏳ Pending: {pending} | 🚩 AI-flagged: {flagged}")


def segments_to_display(segments: list[dict]) -> list[list]:
    """Convert to list-of-lists for gr.Dataframe."""
    rows = []
    for s in segments:
        flag = "🚩" if s["needs_review"] else ""
        rows.append([
            s["id"],
            f"{s['start']:.1f}s–{s['end']:.1f}s",
            s["source_text"],
            s["translated_text"],
            s["corrected_text"],
            f"{s['score']:.2f}",
            ", ".join(s["flags"]) if s["flags"] else "",
            flag,
            s["decision"],
        ])
    return rows
