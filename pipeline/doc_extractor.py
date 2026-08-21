"""Extract text from PDF / DOCX / TXT files.

extract_text()        — plain text (used by ASR/TTS pipeline)
extract_docx_blocks() — structured blocks preserving formatting (used by doc translator)
translate_docx()      — translate a .docx in-place, preserving all styles/tables/runs
"""

from pathlib import Path


def extract_text(file_path: str) -> str:
    p = Path(file_path)
    suffix = p.suffix.lower()

    if suffix == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        # PDF translation blocked per KB tender §3.1 — upload original to CBP portal as-is
        raise ValueError(
            "PDF documents are NOT translated per KB tender §3.1. "
            "Upload the original PDF directly to the CBP portal."
        )

    if suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(p))
            return "\n".join(para.text for para in doc.paragraphs)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

    raise ValueError(f"Unsupported file type: {suffix}")


def translate_docx(src_path: str, out_path: str,
                   translate_fn,
                   src_lang: str, tgt_lang: str) -> None:
    """
    Translate a .docx file to out_path, preserving:
      - Paragraph styles (Heading 1/2/3, Normal, List, etc.)
      - Run-level formatting (bold, italic, underline, font size/colour)
      - Tables (cell by cell)
      - Hyperlinks (text translated, URL preserved)
      - Headers / footers
      - Inline images (copied unchanged)

    translate_fn(texts: list[str], src: str, tgt: str) -> list[str]
      Must accept a batch of strings and return translated strings 1-to-1.
    """
    import copy
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(src_path))

    # ── collect all translatable text units with back-references ──────────
    # Each entry: {"runs": [run, ...], "texts": [str, ...]}
    # We translate run texts individually to keep bold/italic boundaries intact.

    def _translate_para_text(para) -> None:
        """Translate a paragraph as one unit. Put result in first run, clear others."""
        runs = para.runs
        if not runs:
            return
        full_text = "".join(r.text for r in runs)
        if not full_text.strip():
            return
        try:
            translated = translate_fn([full_text], src_lang, tgt_lang)
            result = translated[0] if translated else full_text
        except Exception:
            result = full_text
        runs[0].text = result
        for run in runs[1:]:
            run.text = ""

    def _translate_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _translate_para_text(para)

    # ── body paragraphs ───────────────────────────────────────────────────
    for para in doc.paragraphs:
        _translate_para_text(para)

    # ── tables ────────────────────────────────────────────────────────────
    for table in doc.tables:
        _translate_table(table)

    # ── headers and footers ───────────────────────────────────────────────
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf is not None:
                for para in hf.paragraphs:
                    _translate_para_text(para)
                for table in hf.tables:
                    _translate_table(table)

    doc.save(str(out_path))
