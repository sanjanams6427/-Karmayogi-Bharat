# ============================================================
# Correction Cycle Tracker — KB Tender RFB IN-KBL-543730-NC-RFB
#
# Tender requirements:
#   - Corrections resubmitted within 5 calendar days of feedback
#   - Delay penalty: 0.5% per week (or part thereof) beyond deadline
#   - Correction & Closure Report required per Deliverable 4.5.iv
#
# Ticket lifecycle:  open → in_progress → closed
# Storage: correction_tickets.jsonl  (one JSON object per line, append-only)
# ============================================================

import json
import datetime
from pathlib import Path

TICKETS_FILE = Path(__file__).parent.parent / "translation_memory" / "correction_tickets.jsonl"
DEADLINE_DAYS = 5          # KB tender: 5 calendar days
PENALTY_PCT_PER_WEEK = 0.5  # 0.5% per week (or part thereof)


# ── Helpers ───────────────────────────────────────────────────

def _now_iso() -> str:
    return datetime.datetime.now().isoformat(timespec="seconds")


def _days_overdue(deadline_iso: str) -> int:
    """Return calendar days past deadline (0 if not yet overdue)."""
    deadline = datetime.datetime.fromisoformat(deadline_iso)
    delta = datetime.datetime.now() - deadline
    return max(0, delta.days)


def _weeks_overdue(days_overdue: int) -> int:
    """Ceiling-divide days into weeks (1 day late = 1 week penalty)."""
    if days_overdue <= 0:
        return 0
    return (days_overdue + 6) // 7   # ceil(days / 7)


def penalty_pct(deadline_iso: str, closed_at_iso: str | None = None) -> float:
    """
    Compute the delay penalty percentage.
    If closed_at_iso is given, measure from deadline to closure date.
    Otherwise measure from deadline to now (for open tickets).
    """
    deadline = datetime.datetime.fromisoformat(deadline_iso)
    reference = (datetime.datetime.fromisoformat(closed_at_iso)
                 if closed_at_iso else datetime.datetime.now())
    days_late = max(0, (reference - deadline).days)
    return _weeks_overdue(days_late) * PENALTY_PCT_PER_WEEK


def days_remaining(deadline_iso: str) -> int:
    """Days left until deadline (negative = overdue)."""
    deadline = datetime.datetime.fromisoformat(deadline_iso)
    return (deadline - datetime.datetime.now()).days


# ── Ticket CRUD ───────────────────────────────────────────────

def _load_all() -> list[dict]:
    if not TICKETS_FILE.exists():
        return []
    tickets = []
    for line in TICKETS_FILE.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line:
            try:
                tickets.append(json.loads(line))
            except json.JSONDecodeError:
                pass
    return tickets


def _save_all(tickets: list[dict]) -> None:
    TICKETS_FILE.parent.mkdir(parents=True, exist_ok=True)
    TICKETS_FILE.write_text(
        "\n".join(json.dumps(t, ensure_ascii=False) for t in tickets) + "\n",
        encoding="utf-8",
    )


def raise_ticket(
    course_id: str,
    tgt_lang: str,
    feedback: str,
    raised_by: str = "KB Verification Agency",
    feedback_date: str | None = None,
) -> dict:
    """
    Create a new correction ticket.
    deadline = feedback_date + 5 calendar days.
    Returns the new ticket dict.
    """
    fd = feedback_date or _now_iso()
    fd_dt = datetime.datetime.fromisoformat(fd)
    deadline_dt = fd_dt + datetime.timedelta(days=DEADLINE_DAYS)

    tickets = _load_all()
    ticket_id = f"COR-{len(tickets) + 1:04d}"

    ticket = {
        "ticket_id":     ticket_id,
        "course_id":     course_id,
        "tgt_lang":      tgt_lang,
        "feedback":      feedback,
        "raised_by":     raised_by,
        "raised_at":     fd,
        "deadline":      deadline_dt.isoformat(timespec="seconds"),
        "status":        "open",
        "closed_at":     None,
        "closed_by":     None,
        "resolution":    None,
        "penalty_pct":   0.0,
    }
    tickets.append(ticket)
    _save_all(tickets)
    return ticket


def update_ticket_status(ticket_id: str, status: str) -> dict | None:
    """Set status to 'in_progress'. Returns updated ticket or None."""
    tickets = _load_all()
    for t in tickets:
        if t["ticket_id"] == ticket_id:
            t["status"] = status
            _save_all(tickets)
            return t
    return None


def close_ticket(
    ticket_id: str,
    resolution: str,
    closed_by: str = "Translation Agency",
) -> dict | None:
    """
    Close a ticket. Computes final penalty based on closure date vs deadline.
    Returns updated ticket or None if not found.
    """
    tickets = _load_all()
    for t in tickets:
        if t["ticket_id"] == ticket_id:
            if t["status"] == "closed":
                return t
            closed_at = _now_iso()
            t["status"]     = "closed"
            t["closed_at"]  = closed_at
            t["closed_by"]  = closed_by
            t["resolution"] = resolution
            t["penalty_pct"] = penalty_pct(t["deadline"], closed_at)
            _save_all(tickets)
            return t
    return None


# ── Queries ───────────────────────────────────────────────────

def get_all_tickets() -> list[dict]:
    return _load_all()


def get_open_tickets() -> list[dict]:
    return [t for t in _load_all() if t["status"] != "closed"]


def get_tickets_for_course(course_id: str, tgt_lang: str | None = None) -> list[dict]:
    return [
        t for t in _load_all()
        if t["course_id"] == course_id
        and (tgt_lang is None or t["tgt_lang"] == tgt_lang)
    ]


def ticket_summary(tickets: list[dict]) -> dict:
    """Aggregate stats for a list of tickets."""
    open_t    = [t for t in tickets if t["status"] == "open"]
    inprog    = [t for t in tickets if t["status"] == "in_progress"]
    closed    = [t for t in tickets if t["status"] == "closed"]
    overdue   = [t for t in open_t + inprog if days_remaining(t["deadline"]) < 0]
    total_pen = sum(t["penalty_pct"] for t in closed)
    return {
        "total":       len(tickets),
        "open":        len(open_t),
        "in_progress": len(inprog),
        "closed":      len(closed),
        "overdue":     len(overdue),
        "total_penalty_pct": round(total_pen, 2),
    }


# ── Display helpers ───────────────────────────────────────────

def tickets_to_rows(tickets: list[dict]) -> list[list]:
    """Convert tickets to list-of-lists for gr.Dataframe."""
    rows = []
    for t in tickets:
        dr = days_remaining(t["deadline"])
        if t["status"] == "closed":
            deadline_str = f"Closed {t['closed_at'][:10]}"
            pen_str = f"{t['penalty_pct']:.1f}%"
        elif dr < 0:
            deadline_str = f"⚠️ OVERDUE {abs(dr)}d"
            pen_str = f"{penalty_pct(t['deadline']):.1f}% (accruing)"
        else:
            deadline_str = f"⏳ {dr}d left ({t['deadline'][:10]})"
            pen_str = "0%"
        rows.append([
            t["ticket_id"],
            t["course_id"],
            t["tgt_lang"],
            t["status"],
            t["feedback"][:80],
            deadline_str,
            pen_str,
            t.get("resolution", "") or "",
        ])
    return rows


# ── Closure report (DOCX) ─────────────────────────────────────

def export_closure_report(
    tickets: list[dict],
    output_path: str,
    agency_name: str = "Translation Agency",
) -> str:
    """
    Generate Correction & Closure Report (DOCX) per KB Deliverable 4.5.iv.
    Covers all provided tickets (typically filtered by course_id + lang).
    """
    from docx import Document
    from docx.shared import RGBColor

    summary = ticket_summary(tickets)
    doc = Document()
    doc.add_heading("Correction & Closure Report", level=1)
    doc.add_heading("KB iGOT Karmayogi — Translation Agency", level=2)
    doc.add_paragraph(f"RFB No.: IN-KBL-543730-NC-RFB")
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%d %B %Y %H:%M')}")
    doc.add_paragraph(f"Prepared by: {agency_name}")
    doc.add_paragraph("")

    # Summary
    doc.add_heading("1. Summary", level=2)
    st = doc.add_table(rows=6, cols=2)
    st.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Total Correction Tickets",  str(summary["total"])),
        ("Open",                      str(summary["open"])),
        ("In Progress",               str(summary["in_progress"])),
        ("Closed",                    str(summary["closed"])),
        ("Overdue (open/in-progress)",str(summary["overdue"])),
        ("Total Penalty Accrued",     f"{summary['total_penalty_pct']:.2f}%"),
    ]):
        st.rows[i].cells[0].text = k
        st.rows[i].cells[1].text = v
    doc.add_paragraph("")

    # Ticket detail table
    doc.add_heading("2. Ticket Details", level=2)
    dt = doc.add_table(rows=1, cols=7)
    dt.style = "Table Grid"
    for i, h in enumerate(["Ticket ID", "Course ID", "Language",
                            "Feedback", "Deadline", "Status", "Penalty %"]):
        dt.rows[0].cells[i].text = h

    for t in tickets:
        dr = days_remaining(t["deadline"])
        pen = (t["penalty_pct"] if t["status"] == "closed"
               else penalty_pct(t["deadline"]))
        row = dt.add_row().cells
        row[0].text = t["ticket_id"]
        row[1].text = t["course_id"]
        row[2].text = t["tgt_lang"]
        row[3].text = t["feedback"][:120]
        row[4].text = t["deadline"][:10]
        row[5].text = t["status"]
        row[6].text = f"{pen:.1f}%"
        # Colour overdue rows red
        if t["status"] != "closed" and dr < 0:
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph("")

    # Closed tickets — resolution details
    closed = [t for t in tickets if t["status"] == "closed"]
    if closed:
        doc.add_heading("3. Closed Tickets — Resolution Details", level=2)
        for t in closed:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{t['ticket_id']} ({t['course_id']} / {t['tgt_lang']}): ").bold = True
            p.add_run(t.get("resolution") or "—")
            p.add_run(f"  [Closed {t['closed_at'][:10]} by {t['closed_by']},"
                      f" penalty: {t['penalty_pct']:.1f}%]").italic = True
        doc.add_paragraph("")

    # Compliance declaration
    doc.add_heading("4. Compliance Declaration", level=2)
    open_count = summary["open"] + summary["in_progress"]
    if open_count == 0:
        doc.add_paragraph(
            "All correction tickets have been resolved and closed. "
            "Content meets KB quality standards per RFB IN-KBL-543730-NC-RFB."
        )
    else:
        doc.add_paragraph(
            f"⚠️  {open_count} ticket(s) remain open. "
            "Corrections must be resubmitted within 5 calendar days of feedback "
            "per KB tender §5.1B. Delay penalty: 0.5% per week."
        )
    doc.add_paragraph(f"Submitted by: {agency_name}")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("Signature: ____________________")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
